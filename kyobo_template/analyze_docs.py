"""교보 템플릿 문서 분석 스크립트"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

TEMPLATE_DIR = r"c:\Users\aiueo\projects\quant-alpha\kyobo_template"

def analyze_document(filepath):
    """문서의 구조, 스타일, 시각 요소를 상세 분석"""
    doc = Document(filepath)
    filename = os.path.basename(filepath)
    
    result = {
        "filename": filename,
        "sections": [],
        "styles_used": {},
        "tables": [],
        "images": 0,
        "total_paragraphs": len(doc.paragraphs),
        "fonts_used": set(),
        "colors_used": set(),
        "heading_structure": [],
        "visual_elements": [],
    }
    
    # 1. 문단 분석
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else "None"
        
        # 스타일 카운트
        result["styles_used"][style_name] = result["styles_used"].get(style_name, 0) + 1
        
        # 폰트 분석
        for run in para.runs:
            if run.font.name:
                result["fonts_used"].add(run.font.name)
            if run.font.size:
                result["fonts_used"].add(f"size:{run.font.size.pt}pt")
            if run.font.color and run.font.color.rgb:
                result["colors_used"].add(str(run.font.color.rgb))
            if run.font.bold:
                result["fonts_used"].add("BOLD")
        
        # 헤딩 구조
        if "Heading" in style_name or "heading" in style_name or "제목" in style_name:
            result["heading_structure"].append({
                "level": style_name,
                "text": para.text[:100],
                "index": i
            })
        
        # 빈 문단이 아닌 경우만 섹션에 추가 (첫 100글자)
        if para.text.strip():
            result["sections"].append({
                "index": i,
                "style": style_name,
                "text": para.text[:150],
                "alignment": str(para.alignment) if para.alignment else "None"
            })
    
    # 2. 테이블 분석
    for t_idx, table in enumerate(doc.tables):
        table_info = {
            "table_index": t_idx,
            "rows": len(table.rows),
            "cols": len(table.columns),
            "header_row": [],
            "sample_data": []
        }
        
        # 첫 행 (헤더)
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                table_info["header_row"].append(cell.text[:50])
        
        # 데이터 샘플 (최대 3행)
        for row_idx, row in enumerate(table.rows[1:4], 1):
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text[:50])
            table_info["sample_data"].append(row_data)
        
        result["tables"].append(table_info)
    
    # 3. 이미지/도형 분석
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1
    result["images"] = image_count
    
    # 4. 섹션(페이지 설정) 분석
    page_info = []
    for section in doc.sections:
        page_info.append({
            "page_width": str(section.page_width),
            "page_height": str(section.page_height),
            "left_margin": str(section.left_margin),
            "right_margin": str(section.right_margin),
            "top_margin": str(section.top_margin),
            "bottom_margin": str(section.bottom_margin),
        })
    result["page_settings"] = page_info
    
    # set을 list로 변환
    result["fonts_used"] = sorted(list(result["fonts_used"]))
    result["colors_used"] = sorted(list(result["colors_used"]))
    
    return result


def print_analysis(analysis):
    """분석 결과를 보기 좋게 출력"""
    print(f"\n{'='*80}")
    print(f"📄 {analysis['filename']}")
    print(f"{'='*80}")
    
    print(f"\n📊 기본 정보:")
    print(f"  - 총 문단 수: {analysis['total_paragraphs']}")
    print(f"  - 테이블 수: {len(analysis['tables'])}")
    print(f"  - 이미지 수: {analysis['images']}")
    
    print(f"\n🎨 폰트/스타일:")
    print(f"  - 사용 폰트: {analysis['fonts_used']}")
    print(f"  - 사용 색상(RGB): {analysis['colors_used']}")
    
    print(f"\n📑 스타일 분포:")
    for style, count in sorted(analysis['styles_used'].items(), key=lambda x: -x[1]):
        print(f"  - {style}: {count}회")
    
    print(f"\n📋 헤딩 구조 (목차):")
    for h in analysis['heading_structure']:
        print(f"  [{h['level']}] {h['text']}")
    
    print(f"\n📝 주요 내용 (첫 30개 문단):")
    for s in analysis['sections'][:30]:
        print(f"  [{s['style']}] {s['text'][:100]}")
    
    print(f"\n📊 테이블 상세:")
    for t in analysis['tables']:
        print(f"  Table {t['table_index']}: {t['rows']}행 x {t['cols']}열")
        print(f"    헤더: {t['header_row']}")
        for row in t['sample_data'][:2]:
            print(f"    데이터: {row}")
    
    print(f"\n📐 페이지 설정:")
    for p in analysis['page_settings']:
        print(f"  {p}")


if __name__ == "__main__":
    files = [f for f in os.listdir(TEMPLATE_DIR) if f.endswith('.docx')]
    
    all_analyses = []
    for f in sorted(files):
        filepath = os.path.join(TEMPLATE_DIR, f)
        try:
            analysis = analyze_document(filepath)
            all_analyses.append(analysis)
            print_analysis(analysis)
        except Exception as e:
            print(f"\n❌ {f} 분석 실패: {e}")
    
    # JSON으로도 저장
    output_path = os.path.join(TEMPLATE_DIR, "_analysis_result.json")
    for a in all_analyses:
        a["fonts_used"] = list(a["fonts_used"])
        a["colors_used"] = list(a["colors_used"])
    
    with open(output_path, "w", encoding="utf-8") as fp:
        json.dump(all_analyses, fp, ensure_ascii=False, indent=2)
    
    print(f"\n\n✅ 분석 결과 저장: {output_path}")
